#!/usr/bin/env python3
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_horizontal_rule(doc):
    """Add a horizontal rule (border) between sections."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'C7D2FE')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def set_heading_color(run, level):
    if level == 1:
        run.font.color.rgb = RGBColor(0x31, 0x2E, 0x81)   # indigo-900
        run.font.size = Pt(22)
    elif level == 2:
        run.font.color.rgb = RGBColor(0x43, 0x38, 0xCA)   # indigo-700
        run.font.size = Pt(16)
    elif level == 3:
        run.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)   # blue-700
        run.font.size = Pt(13)

def md_to_docx(md_path, docx_path, title):
    doc = Document()

    # Set page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # Cover block
    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cover.add_run("B20 EXCHANGE")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x43, 0x38, 0xCA)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    srun = sub.add_run(title)
    srun.italic = True
    srun.font.size = Pt(13)
    srun.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    add_horizontal_rule(doc)
    doc.add_paragraph()  # spacing

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_code_block = False
    code_lines = []

    for raw_line in lines:
        line = raw_line.rstrip('\n')

        # Code block handling
        if line.strip().startswith('```'):
            if in_code_block:
                # Close code block — add as styled paragraph
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph()
                p.style = 'No Spacing'
                run = p.add_run(code_text)
                run.font.name = 'Courier New'
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
                in_code_block = False
                code_lines = []
            else:
                in_code_block = True
            continue

        if in_code_block:
            code_lines.append(line)
            continue

        # Skip pure horizontal rules (---)
        if re.match(r'^-{3,}$', line.strip()):
            add_horizontal_rule(doc)
            continue

        # H1
        if line.startswith('# ') and not line.startswith('## '):
            p = doc.add_heading('', level=1)
            run = p.add_run(line[2:].strip())
            run.bold = True
            set_heading_color(run, 1)
            continue

        # H2
        if line.startswith('## '):
            doc.add_paragraph()
            p = doc.add_heading('', level=2)
            run = p.add_run(line[3:].strip())
            run.bold = True
            set_heading_color(run, 2)
            continue

        # H3
        if line.startswith('### '):
            p = doc.add_heading('', level=3)
            run = p.add_run(line[4:].strip())
            run.bold = True
            set_heading_color(run, 3)
            continue

        # Bullet points (* or -)
        if re.match(r'^\s{0,4}[\*\-]\s+', line):
            indent = len(line) - len(line.lstrip())
            text = re.sub(r'^\s*[\*\-]\s+', '', line)
            # Strip bold (**text**) and inline code (`text`)
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text = re.sub(r'`(.*?)`', r'\1', text)
            style = 'List Bullet 2' if indent >= 4 else 'List Bullet'
            p = doc.add_paragraph(text, style=style)
            continue

        # Numbered lists
        if re.match(r'^\d+\.\s+', line):
            text = re.sub(r'^\d+\.\s+', '', line)
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text = re.sub(r'`(.*?)`', r'\1', text)
            doc.add_paragraph(text, style='List Number')
            continue

        # Empty line
        if not line.strip():
            doc.add_paragraph()
            continue

        # Normal paragraph — strip inline markdown
        clean = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
        clean = re.sub(r'\*(.*?)\*', r'\1', clean)
        clean = re.sub(r'`(.*?)`', r'\1', clean)
        p = doc.add_paragraph(clean)
        p.style.font.size = Pt(10)

    # Footer
    add_horizontal_rule(doc)
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer_p.add_run("© B20 Exchange · Nexus Nuera Admin · Confidential Document")
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
    fr.italic = True

    doc.save(docx_path)
    print(f"✅ Saved: {docx_path}")

md_to_docx(
    '/Users/purusothaman/Desktop/myfinal/B20_Technical_Blueprint.md',
    '/Users/purusothaman/Desktop/myfinal/frontend/public/B20_Technical_Blueprint.docx',
    'Institutional Technical Blueprint'
)

md_to_docx(
    '/Users/purusothaman/Desktop/myfinal/B20_MASTER_AUDIT_DOCUMENT.md',
    '/Users/purusothaman/Desktop/myfinal/frontend/public/B20_Master_Audit_Document.docx',
    'Omnibus Ecosystem Audit & Technical Reference'
)

print("✅ All documents converted successfully.")
